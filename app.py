import docx
from docx import Document
import gradio as gr
import openai
openai.api_key="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" # API key from OpenAI account (personal and unique)

# A function to translate the text
def summerize_text(input_file):
    
    doc=Document(input_file)

    text=[]

    for para in doc.paragraphs:
        # print(para.text)
        text.append(para.text)

    prompt = f"Summerize the following {text}"

    summerized = openai.ChatCompletion.create(

        model="gpt-3.5-turbo", # ChatGPT model

        # The system message helps set the behavior of the assistant.
        # Assistant messages store previous assistant responses, but can also be written by you to give examples of desired behavior.
        # The user messages provide requests or comments for the assistant to respond to.
        messages=[{"role": "system", "content": "You are a helpful assistant to summerize the text."},
                  {"role": "assistant","content":"Text summerizer"},
                  {"role": "user", "content": prompt}],

        # OpenAI models are non-deterministic, meaning that identical inputs can yield different outputs.
        # Setting temperature to 0 will make the outputs mostly deterministic, but a small amount of variability may remain.
        temperature=0.5,
    ) # End of ChatCompletion

    summary = summerized.choices[0].message.content.strip()

    doc.add_paragraph(summary)
    doc.save(r"C:\Data Science\Assignments\New Microsoft Word Document.docx") # Location where the file will be saved

    return doc

iface = gr.Interface(fn=summerize_text,inputs=[gr.File(label="Upload a DOCX file")],outputs=gr.File(label="Download summerized DOCX"),
    title="Script Summerisation App",
    description="Summerize the script DOCX file.",
)

iface.launch()